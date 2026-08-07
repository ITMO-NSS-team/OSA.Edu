from __future__ import annotations

import json
import re
from pathlib import Path
from statistics import median
from typing import Any

from .config import EXTRACTED_DIR
from .document.analyze import analyze_document
from .pdf_blocks import build_pdf_pages
from .util import normalize_text


def extract_document(path: str | Path) -> dict[str, Any]:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == '.pdf':
        return _extract_pdf(path)
    if suffix == '.docx':
        return _extract_docx(path)
    raise ValueError('Поддерживаются только PDF и DOCX.')


def _extract_pdf(path: Path) -> dict[str, Any]:
    import pymupdf

    document = pymupdf.open(path)
    try:
        if document.needs_pass:
            raise RuntimeError('PDF защищён паролем и не может быть прочитан без пароля.')
        pages, style_stats, empty_pages = build_pdf_pages(document)
    finally:
        document.close()

    warnings: list[str] = []
    if empty_pages:
        warnings.append('На страницах без текстового слоя потребуется OCR: ' + ', '.join(map(str, empty_pages)) + '.')
    text = '\n\n'.join(f"<<<PAGE {page['number']}>>>\n{page['text']}" for page in pages)
    return analyze_document(
        text,
        pages,
        'pdf',
        warnings or ([] if pages else ['Не удалось определить страницы PDF.']),
        style_stats=style_stats,
    )

def _extract_docx(path: Path) -> dict[str, Any]:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(path)
    body = doc.element.body
    paragraphs = {p._element: p for p in doc.paragraphs}
    tables = {t._element: t for t in doc.tables}
    raw_blocks: list[dict[str, Any]] = []
    font_stats: dict[str, int] = {}
    size_stats: dict[str, int] = {}
    alignment_stats: dict[str, int] = {}

    def heading_level(paragraph, element) -> int | None:
        style = paragraph.style.name if paragraph.style is not None else ''
        match = re.search(r'(?:heading|заголовок)\s*(\d)', style, re.I)
        if match:
            return int(match.group(1))
        if style.strip().lower() in {'title', 'название'}:
            return 0
        ppr = element.find(qn('w:pPr'))
        outline = ppr.find(qn('w:outlineLvl')) if ppr is not None else None
        value = outline.get(qn('w:val')) if outline is not None else None
        return int(value) + 1 if value and value.isdigit() and int(value) < 9 else None

    for source_order, child in enumerate(body.iterchildren()):
        if child.tag == qn('w:p') and child in paragraphs:
            paragraph = paragraphs[child]
            text = normalize_text(paragraph.text)
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style is not None else '') or ''
            level = heading_level(paragraph, child)
            ppr = child.find(qn('w:pPr'))
            numbered = ppr is not None and ppr.find(qn('w:numPr')) is not None
            runs = [r for r in paragraph.runs if r.text.strip()]
            total = sum(len(r.text) for r in runs)
            bold = total > 0 and sum(len(r.text) for r in runs if r.bold) / total >= .6
            italic = total > 0 and sum(len(r.text) for r in runs if r.italic) / total >= .6
            dominant_fonts: dict[str, int] = {}
            sizes: list[float] = []
            for run in runs:
                length = max(1, len(run.text.strip()))
                font = run.font.name or (paragraph.style.font.name if paragraph.style is not None else None)
                if font:
                    dominant_fonts[font] = dominant_fonts.get(font, 0) + length
                    font_stats[font] = font_stats.get(font, 0) + length
                size = run.font.size.pt if run.font.size is not None else None
                if size:
                    sizes.extend([float(size)] * length)
                    key = f'{float(size):g}'
                    size_stats[key] = size_stats.get(key, 0) + length
            alignment = str(paragraph.paragraph_format.alignment or 'INHERIT').split('.')[-1]
            alignment_stats[alignment] = alignment_stats.get(alignment, 0) + 1
            page_break_before = False
            if ppr is not None and ppr.find(qn('w:pageBreakBefore')) is not None:
                page_break_before = True
            raw_blocks.append({
                'text': text,
                'style': style,
                'level': level,
                'listitem': bool(numbered) or 'list' in style.lower() or 'список' in style.lower(),
                'fontName': max(dominant_fonts, key=dominant_fonts.get) if dominant_fonts else None,
                'fontSize': round(median(sizes), 2) if sizes else None,
                'bold': bold,
                'italic': italic,
                'alignment': alignment,
                'pageBreakBefore': page_break_before,
                'sourceOrder': source_order,
            })
        elif child.tag == qn('w:tbl') and child in tables:
            table = tables[child]
            rows: list[str] = []
            for row in table.rows:
                cells = [normalize_text(cell.text).replace('\n', ' ') for cell in row.cells]
                rows.append(' | '.join(cells))
            table_text = '\n'.join(x for x in rows if x.strip(' |'))
            if table_text:
                raw_blocks.append({'text': table_text, 'type': 'table', 'style': 'Table', 'sourceOrder': source_order})

    text = normalize_text('\n\n'.join(str(x.get('text') or '') for x in raw_blocks))
    style_stats = {'fonts': font_stats, 'sizes': size_stats, 'alignment': alignment_stats}
    return analyze_document(
        text,
        [],
        'docx',
        ['DOCX не содержит надёжной привязки к страницам. Для финальной проверки вёрстки загрузите также PDF.'],
        raw_blocks=raw_blocks,
        style_stats=style_stats,
    )


def extracted_path(job_id: str) -> Path:
    return EXTRACTED_DIR / f'{job_id}.json'


def save_extracted(job_id: str, document: dict[str, Any]) -> str:
    path = extracted_path(job_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(document, ensure_ascii=False, indent=2), encoding='utf-8')
    return str(path)


def read_extracted(path: str | Path) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding='utf-8'))
